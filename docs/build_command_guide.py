from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Guia_comandos_controladores_Crazyflie.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
NAVY = "16324F"
MUTED = "667085"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
RED = "9B1C1C"
GOLD = "7A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def font_run(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    font_run(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "TESIS · Crazyflie 2.1 · Guía de operación"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font_run(header.runs[0], size=9, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GUÍA DE COMANDOS")
    font_run(run, size=28, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Controladores Crazyflie con Robotat y Flow deck")
    font_run(run, size=16, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    run = p.add_run("Referencia rápida para diagnóstico, vuelo, teclado, cámara y gestos")
    font_run(run, size=11, color=MUTED, italic=True)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = (("Plataforma", "Crazyflie 2.1"), ("Proyecto", "Tesis · ecosistema Robotat"), ("Actualización", "Agosto de 2026"))
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT)
        font_run(row.cells[0].paragraphs[0].add_run(label), bold=True, color=NAVY)
        font_run(row.cells[1].paragraphs[0].add_run(value))
    set_table_geometry(table, (2700, 6660))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento operativo: verifica siempre el sistema de posicionamiento antes de despegar.")
    font_run(run, size=10, color=RED, bold=True)
    doc.add_page_break()


def add_command(doc, command, description=None):
    command = command.replace("\\\\", "\\")
    if description:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(description)
        font_run(run, bold=True, color=NAVY)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(command)
    font_run(run, name="Consolas", size=9.3, color="111827")
    set_table_geometry(table, (9360,))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_note(doc, label, text, color=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_shading(table.cell(0, 0), "FFF8E8" if color == GOLD else "FDECEC")
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run(f"{label}: ")
    font_run(r, bold=True, color=color)
    font_run(p.add_run(text), color="222222")
    set_table_geometry(table, (9360,))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_key_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], LIGHT)
        font_run(table.rows[0].cells[index].paragraphs[0].add_run(header), bold=True, color=NAVY)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            font_run(cells[index].paragraphs[0].add_run(value), name="Consolas" if index == 0 else "Calibri", size=10)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, widths)
    return table


def build():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Inicio rápido", level=1)
    p = doc.add_paragraph("Todos los comandos deben ejecutarse desde la carpeta raíz del proyecto: ")
    font_run(p.add_run(str(ROOT)), name="Consolas", size=9.5, color=DARK)
    add_command(doc, r".\.venv\Scripts\python.exe --version", "Comprobar el entorno virtual")
    add_note(doc, "Importante", "Usa siempre .\\.venv\\Scripts\\python.exe y no el comando python del sistema.")

    doc.add_heading("2. Seguridad antes de volar", level=1)
    for label, text in (
        ("Área", "Despeja personas, cables y objetos; usa buena iluminación y suelo mate con textura."),
        ("Flow deck", "No despegues si no aparece 'Flow deck detectado' y 'Estimador estable'."),
        ("Robotat", "No despegues si MoCap está desactualizado o el EKF no está alineado."),
        ("Aterrizaje", "Usa el botón ATERRIZAR o q en la cámara para un cierre normal."),
        ("Emergencia", "Q o el puño cerrado cortan motores; el dron caerá. Úsalos solo ante peligro."),
    ):
        p = doc.add_paragraph()
        font_run(p.add_run(f"{label}. "), bold=True, color=NAVY)
        font_run(p.add_run(text))

    doc.add_heading("3. Diagnóstico de Flow deck", level=1)
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\diagnostico_flowdeck_dos_drones.py", "Comprobar ambos drones, uno después del otro, sin motores")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\diagnostico_flowdeck_dos_drones.py --tiempo 10", "Ampliar el muestreo a 10 segundos por dron")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\diagnostico_flowdeck_dos_drones.py --radio1 2B1D933FCC --radio2 9DD2507072", "Elegir las antenas manualmente")
    add_note(doc, "Resultado esperado", "deck.bcFlow2=1, cambio de range.zrange y muestras no nulas de motion.deltaX/Y.")

    doc.add_heading("4. Flow deck · un dron", level=1)
    doc.add_heading("4.1 Hover básico del dron 1", level=2)
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\single_drone\flowdeck\hover_flowdeck_dron1.py")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\single_drone\flowdeck\hover_flowdeck_dron1.py --altura 0.45 --tiempo 10", "Definir altura y duración")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\single_drone\flowdeck\hover_flowdeck_dron1.py --radio 9DD2507072", "Seleccionar una Crazyradio")
    add_key_table(doc, ("Entrada", "Acción"), (("Enter", "Autorizar el despegue"), ("Q", "Cancelar en tierra o cortar motores durante el vuelo"), ("Ctrl+C", "Aterrizaje/cierre normal")), (2500, 6860))

    doc.add_heading("4.2 Panel de teclado del dron 1", level=2)
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\single_drone\flowdeck\panel_control_flowdeck_dron1.py")
    add_key_table(doc, ("Tecla", "Acción"), (("W / S", "Adelante / atrás"), ("A / D", "Izquierda / derecha"), ("Espacio / Shift", "Subir / bajar"), ("Q", "Corte inmediato de motores"), ("Ctrl+C", "Aterrizar y cerrar")), (2500, 6860))

    doc.add_heading("4.3 Cámara del dron 1", level=2)
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\single_drone\camera\control_camara_flowdeck_dron1.py")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\single_drone\camera\control_camara_flowdeck_dron1.py --camera 1", "Usar la cámara secundaria")

    doc.add_heading("5. Flow deck · dos drones", level=1)
    add_note(doc, "Requisito", "Para volar ambos al mismo tiempo se necesitan dos Crazyradio y dos Flow decks funcionales.")
    doc.add_heading("5.1 Panel dual", level=2)
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\panel_control_flowdeck_dos_drones.py")
    add_key_table(doc, ("Dron 1", "Dron 2", "Acción"), (("W / S", "↑ / ↓", "Adelante / atrás"), ("A / D", "← / →", "Izquierda / derecha"), ("Espacio / Shift", "Page Up / Page Down", "Subir / bajar"), ("Q", "Q", "Emergencia de ambos")), (2800, 3000, 3560))

    doc.add_heading("5.2 Cámara dual", level=2)
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_camara_flowdeck_dos_drones.py --target hands", "Mano izquierda=Dron 1; mano derecha=Dron 2")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_camara_flowdeck_dos_drones.py --target drone1", "Controlar solo el dron 1")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_camara_flowdeck_dos_drones.py --target drone2", "Controlar solo el dron 2")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_camara_flowdeck_dos_drones.py --target both", "Una mano controla ambos")

    doc.add_heading("6. Gestos de mano", level=1)
    add_key_table(doc, ("Gesto", "Comando"), (("Índice + medio hacia arriba", "Despegar"), ("Índice + medio hacia abajo", "Aterrizar"), ("Pulgar + índice", "Adelante"), ("Pulgar + meñique", "Atrás"), ("Solo meñique", "Izquierda"), ("Solo pulgar", "Derecha"), ("Solo índice arriba / abajo", "Subir / bajar"), ("Sin mano / reposo", "Hover"), ("Puño durante 0.6 s", "Emergencia"), ("q en la ventana", "Aterrizar y salir")), (5000, 4360))

    doc.add_heading("7. Controladores con Robotat", level=1)
    doc.add_heading("7.1 Un dron", level=2)
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\single_drone\buttons\control_dron_individual_interfaz.py --drone 1", "Panel individual del dron 1")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\single_drone\buttons\control_dron_individual_interfaz.py --drone 2", "Panel individual del dron 2")
    doc.add_heading("7.2 Dos drones por botones", level=2)
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_dos_drones_botones_lowlevel.py", "Control low-level con MoCap")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_dos_drones_cruz_botones.py", "Control high-level por botones")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_dos_drones_cruz_botones.py --dry-run", "Simulación sin radios ni motores")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_dos_drones_cruz_botones.py --single drone2", "Habilitar solo el dron 2")
    doc.add_heading("7.3 Cámara con Robotat", level=2)
    add_command(doc, r".\.venv\Scripts\python.exe .\control_dos_drones_camara.py", "Dos drones: mano izquierda/derecha")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_dos_drones_cruz_camara_multiprocessing.py --target drone2", "Cámara para un único dron")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\control_dos_drones_cruz_camara_multiprocessing.py --dry-run", "Prueba de cámara sin activar hardware")

    doc.add_heading("8. Marker joystick con Robotat", level=1)
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\joystick\discover_marker_id.py --id 64 --show-all", "Encontrar el tópico del marker")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\joystick\marker_orientation_check.py --marker-topic mocap/all --marker-id 64", "Verificar orientación sin conectar el dron")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\joystick\control_with_marker.py --marker-topic mocap/all --marker-id 64", "Iniciar el control de vuelo")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\joystick\analyze_marker_session.py", "Regenerar gráficas de la sesión")

    doc.add_heading("9. Pruebas y análisis", level=1)
    add_command(doc, r".\.venv\Scripts\python.exe .\external\gesture_detection\main_hands.py", "Probar únicamente el reconocimiento de gestos")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\prueba_estabilidad_dos_drones_lowlevel.py", "Prueba de estabilidad de dos drones con Robotat")
    add_command(doc, r".\.venv\Scripts\python.exe .\controllers\two_drones\analizar_sesion_dos_drones.py", "Analizar el CSV más reciente")

    doc.add_heading("10. Cierre y solución rápida de problemas", level=1)
    add_key_table(doc, ("Síntoma", "Acción recomendada"), (("deck.bcFlow2=0", "Apagar, desconectar batería, revisar pines y probar el deck en el otro dron."), ("El dron deriva", "Usar suelo mate con textura, buena luz y revisar hélices/motores."), ("Una tecla queda activa", "Soltar teclas, hacer clic en la ventana; el deadman manda hover en 0.8 s."), ("La cámara no abre", "Cerrar otras aplicaciones o ejecutar con --camera 1."), ("Robotat no actualiza", "No despegar; revisar MQTT, tópico e ID del rigid body."), ("El programa no cierra", "Usar aterrizaje normal y Ctrl+C; Q solo ante emergencia.")), (2800, 6560))
    add_note(doc, "Regla final", "Después de una caída o corte de emergencia, inspecciona hélices, motores, deck y conectores antes del siguiente vuelo.", color=RED)

    doc.core_properties.title = "Guía de comandos de controladores Crazyflie"
    doc.core_properties.subject = "Robotat, Flow deck, teclado, cámara y diagnóstico"
    doc.core_properties.author = "Proyecto de tesis"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
